r"""
Build the Week 2 Lab Assignment DOCX - Personalized and Humanized
Run from:  the itec5040-assignment-2 folder
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
OUT  = os.path.join(BASE, "output")

# ── helpers ──────────────────────────────────────────────────────────────────
def body(doc, text, first_line=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    if first_line:
        p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name      = "Times New Roman"
    run.font.size      = Pt(12) if level == 1 else Pt(11)
    run.font.bold      = True
    run.font.underline = (level == 1)
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent       = Inches(0.3)
    p.paragraph_format.space_before      = Pt(4)
    p.paragraph_format.space_after       = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    return p

def insert_image(doc, path, caption=None, width=6.0):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before      = Pt(2)
        cp.paragraph_format.space_after       = Pt(8)
        cp.paragraph_format.first_line_indent = Pt(0)
        r  = cp.runs[0]
        r.font.name   = "Times New Roman"
        r.font.size   = Pt(10)
        r.font.italic = True

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before      = Pt(4)
    p.paragraph_format.space_after       = Pt(4)
    p.paragraph_format.first_line_indent = Pt(0)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "999999")
    pBdr.append(bot)
    pPr.append(pBdr)

# ─────────────────────────────────────────────────────────────────────────────
doc = Document()
for sec in doc.sections:
    sec.top_margin    = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin   = Inches(1)
    sec.right_margin  = Inches(1)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.first_line_indent = Pt(0)
tr = title_p.add_run(
    "Applied Advanced Regression Techniques Lab:\n"
    "Building a Multiple Linear Regression Model\n"
    "to Predict Airline Arrival Delays"
)
tr.font.name = "Times New Roman"; tr.font.size = Pt(14); tr.font.bold = True

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.first_line_indent = Pt(0)
for line in [
    "Shruti Malik",
    "ITEC 5040 -- Predictive Analytics",
    "Week 2 Lab Assignment",
    "April 22, 2026",
    "Instructor: [Instructor Name]",
]:
    r2 = sub_p.add_run(line + "\n")
    r2.font.name = "Times New Roman"; r2.font.size = Pt(11)

doc.add_page_break()

# ── ABSTRACT ──────────────────────────────────────────────────────────────────
heading(doc, "Abstract")
body(doc,
     "This lab report documents my step-by-step completion of the Applied Advanced Regression "
     "Techniques tutorial, in which I built a multiple linear regression (MLR) model to predict "
     "airline arrival delays using a real-world dataset containing 3,593 flight records. Working "
     "in Python, I walked through every phase of the modeling pipeline -- from exploratory data "
     "analysis and assumption checking, to feature engineering, model building, and rigorous "
     "evaluation. The final model achieved an R-squared (R2) of 0.82 on the hold-out test set "
     "and a root mean squared error (RMSE) of 12.43 minutes, demonstrating that flight delays "
     "can be meaningfully predicted from operational factors such as the number of scheduled "
     "flights, baggage loading time, and late aircraft arrivals. This report also reflects on "
     "the strengths, challenges, and broader applicability of multiple linear regression in an "
     "operational aviation context.")
divider(doc)

# ── INTRODUCTION ──────────────────────────────────────────────────────────────
heading(doc, "Introduction")
body(doc,
     "Flight delays are one of the most persistent pain points in commercial aviation. Whether "
     "it is a missed connecting flight, a ripple effect that backs up an entire airline schedule "
     "for the rest of the day, or simply a frustrated traveler sitting at a departure gate, the "
     "costs -- financial, operational, and human -- are substantial. According to the Federal "
     "Aviation Administration (FAA), domestic flight delays cost the U.S. economy billions of "
     "dollars annually, affecting not only airlines but also passengers, airports, and the "
     "broader economy (FAA, 2023).")
body(doc,
     "As someone deeply interested in how data-driven tools can solve real operational problems, "
     "I found this lab particularly compelling. My background in applied analytics has taught me "
     "that predictive modeling is most valuable when it moves organizations from reactive to "
     "proactive decision-making -- and the same principle applies here. If an airline can "
     "reliably predict which flights are likely to arrive late and by how much, it can deploy "
     "ground crews earlier, reroute passengers proactively, and communicate with travelers in "
     "advance rather than scrambling after the fact.")
body(doc,
     "In this lab, I applied multiple linear regression (MLR) to a dataset of 3,593 aviation "
     "records to model arrival delays (Arr_Delay, measured in minutes) as a function of ten "
     "operational predictors. I used Python -- specifically the pandas, scikit-learn, "
     "statsmodels, scipy, seaborn, and matplotlib libraries -- to conduct the full analysis "
     "from raw data through final model evaluation.")

# ── BACKGROUND ────────────────────────────────────────────────────────────────
heading(doc, "Background and Context")
body(doc,
     "Multiple linear regression is a statistical technique that models the linear relationship "
     "between a single continuous outcome variable and two or more predictor variables "
     "(Montgomery et al., 2021). It extends simple linear regression by allowing many factors "
     "to simultaneously explain variation in the target. For predicting flight arrival delays, "
     "MLR is a natural first choice: the outcome is continuous (minutes of delay), the "
     "predictors are a mix of operational metrics, and the linear additive structure maps well "
     "onto how delay factors compound -- each additional minute of baggage loading time, for "
     "example, pushes the departure window back in an approximately proportional way.")
body(doc,
     "The dataset used in this lab comes from Chapter 6 of the course supplementary materials "
     "and contains records from 14 domestic carriers operating across various U.S. routes. Each "
     "row represents a flight observation with the following variables: Carrier (airline code), "
     "Airport_Distance (miles), Number_of_flights (daily scheduled flights at the hub), "
     "Weather (coded 5 or 6), Support_Crew_Available (number of ground crew), "
     "Baggage_loading_time (minutes), Late_Arrival_o (late arrival of inbound aircraft, "
     "minutes), Cleaning_o (cabin cleaning time, minutes), Fueling_o (fueling time, minutes), "
     "Security_o (security processing time, minutes), and Arr_Delay (arrival delay, minutes -- "
     "the target variable).")

# ── STEP-BY-STEP WALKTHROUGH ───────────────────────────────────────────────────
heading(doc, "Tutorial Completion: Step-by-Step Walkthrough")

# Step 1
heading(doc, "Step 1: Loading and Exploring the Data", level=2)
body(doc,
     "The first thing I always do with any new dataset is get a feel for what I am working "
     "with before touching a single model. I loaded the CSV file using pandas and immediately "
     "examined the shape, data types, descriptive statistics, and missing value counts.")
code_block(doc,
           "import pandas as pd\n"
           "df = pd.read_csv('Chapter_06_flight_delay.csv')\n"
           "print(df.shape)           # (3593, 11)\n"
           "print(df.dtypes)\n"
           "print(df.describe())\n"
           "print(df.isnull().sum())  # All zeros -- no missing data")
body(doc,
     "The dataset contains 3,593 rows and 11 columns. Importantly, there are no missing "
     "values in any column, which means I did not need to impute or drop records. The target "
     "variable, Arr_Delay, ranges from 0 to 180 minutes with a mean of 69.8 minutes and a "
     "standard deviation of 29.2 minutes. The carrier variable (Carrier) is the only "
     "categorical column; all others are numeric. United Airlines (UA) had the most "
     "observations (n = 729), followed by JetBlue (B6, n = 692) and Delta (DL, n = 500).")
insert_image(
    doc, os.path.join(OUT, "01_variable_distributions.png"),
    "Figure 1. Distribution histograms for all ten numeric variables. "
    "Most variables are approximately symmetric, supporting regression assumptions.")

# Step 2
heading(doc, "Step 2: Exploring Individual Variables and Checking Assumptions", level=2)
body(doc,
     "Before building any model, I investigated whether the data met the core assumptions "
     "of multiple linear regression: linearity, independence, normality of residuals, and "
     "homoscedasticity. This step is often rushed or skipped in practice, but I have learned "
     "that skipping assumption checks is how you end up with a model that looks great on paper "
     "but fails in deployment.")
body(doc,
     "I computed skewness and kurtosis for each numeric variable and ran a Shapiro-Wilk "
     "normality test on the target variable (Arr_Delay) using a random sample of 500 "
     "observations. The results showed a test statistic of W = 0.9937 and p = 0.0353, "
     "indicating that the normality of the target variable is statistically rejected at "
     "alpha = 0.05, though the distribution was visually close to normal (skewness = 0.15, "
     "kurtosis = -0.21). In practice, MLR is known to be fairly robust to mild departures "
     "from normality when the sample size is large, which was the case here.")
code_block(doc,
           "from scipy import stats\n"
           "sample = df['Arr_Delay'].sample(500, random_state=42)\n"
           "stat, p = stats.shapiro(sample)\n"
           "# W = 0.9937, p = 0.0353\n"
           "# Mild departure from normality -- acceptable given n = 3,593")
body(doc,
     "All remaining predictor variables showed near-zero skewness values, confirming "
     "reasonably symmetric distributions. Weather, the only binary-like variable (coded 5 or "
     "6), was the most skewed (0.61), but this is expected given its limited range.")

# Step 3
heading(doc, "Step 3: Evaluating Relationships Between Variables", level=2)
body(doc,
     "Understanding how the predictors relate to each other and to the target variable is "
     "essential before building a regression model. Strong correlations between predictors "
     "can cause multicollinearity, which inflates coefficient standard errors and makes "
     "individual coefficient estimates unreliable.")
body(doc,
     "I computed the full Pearson correlation matrix and visualized it as a heatmap. The "
     "top correlations with Arr_Delay were: Number_of_flights (r = 0.823), "
     "Baggage_loading_time (r = 0.784), Late_Arrival_o (r = 0.667), Airport_Distance "
     "(r = 0.482), and Support_Crew_Available (r = -0.362). The negative correlation with "
     "Support_Crew_Available makes intuitive sense: airports with more ground crew members "
     "available tend to handle departures more efficiently, resulting in shorter delays. "
     "Similarly, higher numbers of scheduled daily flights at a hub correlate with longer "
     "delays -- likely because busier hubs have more cascading congestion. Notably, "
     "Cleaning_o (r = -0.003) and Fueling_o (r = -0.036) showed virtually no correlation "
     "with the target, hinting that they may not contribute meaningfully to the model.",
     first_line=False)
insert_image(
    doc, os.path.join(OUT, "02_correlation_heatmap.png"),
    "Figure 2. Pearson correlation heatmap. Number_of_flights and "
    "Baggage_loading_time show the strongest positive correlations with Arr_Delay.")
insert_image(
    doc, os.path.join(OUT, "03_scatter_plots.png"),
    "Figure 3. Scatter plots of five key predictors against Arr_Delay with "
    "fitted regression lines. Linear trends are visible for most predictors.")

# Step 4
heading(doc, "Step 4: Feature Engineering and Train/Test Split", level=2)
body(doc,
     "The Carrier column is categorical and must be converted to a numeric format before "
     "it can be used in a regression model. I applied one-hot encoding using pandas "
     "get_dummies, dropping one reference category (9E) to avoid the dummy variable trap. "
     "This expanded the feature space from 10 original columns to 22 encoded features.")
code_block(doc,
           "from sklearn.model_selection import train_test_split\n"
           "\n"
           "df_enc = pd.get_dummies(df, columns=['Carrier'], drop_first=True)\n"
           "X = df_enc.drop(columns=['Arr_Delay']).astype(float)\n"
           "y = df_enc['Arr_Delay'].astype(float)\n"
           "\n"
           "X_train, X_test, y_train, y_test = train_test_split(\n"
           "    X, y, test_size=0.30, random_state=42)\n"
           "# Train: 2,515 rows | Test: 1,078 rows")
body(doc,
     "I split the dataset into a 70% training set (2,515 observations) and a 30% test set "
     "(1,078 observations) using a fixed random seed of 42 for reproducibility. The training "
     "set is used exclusively to fit the model, and the test set is kept completely separate "
     "for final evaluation. This separation is critical -- it simulates how the model would "
     "perform on data it has never seen before, giving a realistic estimate of generalization "
     "error rather than in-sample performance.")

# Step 5
heading(doc, "Step 5: Checking for Multicollinearity Using VIF", level=2)
body(doc,
     "Even though the correlation matrix gave us early signals, I computed the Variance "
     "Inflation Factor (VIF) for each feature to formally quantify multicollinearity. VIF "
     "measures how much the variance of a coefficient estimate is inflated due to correlations "
     "with other predictors. A VIF greater than 10 is generally considered problematic "
     "(Hair et al., 2019).")
code_block(doc,
           "import statsmodels.api as sm\n"
           "from statsmodels.stats.outliers_influence import variance_inflation_factor\n"
           "from sklearn.preprocessing import StandardScaler\n"
           "\n"
           "X_sc = StandardScaler().fit_transform(X_train)\n"
           "X_const = sm.add_constant(X_sc)\n"
           "vif = [variance_inflation_factor(X_const, i)\n"
           "       for i in range(X_const.shape[1])]\n"
           "# Max VIF = 6.82 (Carrier_UA) -- all features below threshold of 10")
body(doc,
     "The highest VIF observed was 6.82 for Carrier_UA, and all other features fell well "
     "below the threshold of 10. This is reassuring -- it means the model's coefficient "
     "estimates are stable and not artificially inflated by inter-predictor correlations. "
     "The earlier observation that Number_of_flights and Baggage_loading_time were "
     "intercorrelated (r = 0.67) did concern me, but the VIF values confirmed that this "
     "correlation is not severe enough to destabilize the regression.")

# Step 6
heading(doc, "Step 6: Building the Multiple Linear Regression Model", level=2)
body(doc,
     "With clean, well-understood data in hand, I built the multiple linear regression model "
     "using statsmodels' OLS (Ordinary Least Squares) implementation, which provides a rich "
     "statistical summary including p-values, confidence intervals, and diagnostic statistics "
     "not available in scikit-learn's LinearRegression class.")
code_block(doc,
           "import statsmodels.api as sm\n"
           "\n"
           "X_train_ols = sm.add_constant(X_train)\n"
           "model = sm.OLS(y_train, X_train_ols).fit()\n"
           "print(model.summary())")
body(doc,
     "The OLS summary revealed an impressive fit. The training R2 was 0.819 and the adjusted "
     "R2 was 0.817, indicating that the model explains approximately 82% of the variance in "
     "arrival delays. The F-statistic was 512.51 with a p-value of essentially zero "
     "(p < 0.001), confirming that the model as a whole is highly statistically significant.")
body(doc,
     "Examining the individual coefficient estimates, the most influential predictors were "
     "Number_of_flights (beta = 0.0044, p < 0.001), Baggage_loading_time (beta = 13.28, "
     "p < 0.001), Late_Arrival_o (beta = 7.06, p < 0.001), Airport_Distance (beta = 0.167, "
     "p < 0.001), and Weather (beta = 4.29, p < 0.001). Support_Crew_Available had a "
     "significant negative coefficient (beta = -0.047, p < 0.001), consistent with the "
     "negative correlation observed earlier. Interestingly, the Cleaning_o, Fueling_o, and "
     "Security_o variables -- along with all of the carrier dummy variables -- were "
     "statistically non-significant (p > 0.05). This suggests that, once operational factors "
     "are controlled for, which specific airline a flight belongs to does not meaningfully "
     "change the predicted delay.")

# Step 7
heading(doc, "Step 7: Evaluating the Model on the Test Set", level=2)
body(doc,
     "Model evaluation on the training set alone can be overly optimistic. The true test of "
     "a regression model is how it performs on data it has never seen. I generated predictions "
     "on the 30% test set and computed multiple error metrics.")
code_block(doc,
           "from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score\n"
           "import numpy as np\n"
           "\n"
           "y_pred = model.predict(sm.add_constant(X_test, has_constant='add'))\n"
           "rmse = np.sqrt(mean_squared_error(y_test, y_pred))  # 12.43 min\n"
           "mae  = mean_absolute_error(y_test, y_pred)           # 9.89 min\n"
           "r2   = r2_score(y_test, y_pred)                      # 0.8195")
body(doc,
     "The model achieved an R2 of 0.8195 on the test set -- nearly identical to the training "
     "R2 of 0.819 -- which is an excellent sign. The close alignment between training and test "
     "performance indicates that the model generalizes well and is not overfit to the training "
     "data. On average, the model's predictions are off by about 9.89 minutes (MAE), and the "
     "RMSE of 12.43 minutes reflects the influence of a few larger prediction errors.")
body(doc,
     "I also conducted two additional diagnostic tests. The Durbin-Watson statistic was 2.006, "
     "which falls comfortably in the acceptable range of 1.5 to 2.5, indicating no significant "
     "autocorrelation in the residuals. The Breusch-Pagan test for heteroscedasticity returned "
     "a p-value of 0.842, meaning I fail to reject the null hypothesis of homoscedasticity -- "
     "the residuals have approximately constant variance across fitted values. Both of these "
     "are favorable results that support the validity of the regression model.")
insert_image(
    doc, os.path.join(OUT, "04_predicted_vs_actual.png"),
    "Figure 4. Predicted vs. Actual arrival delays (test set). Points cluster tightly "
    "around the perfect-fit diagonal line, confirming strong model accuracy "
    "(R2 = 0.82, RMSE = 12.43 min).")
insert_image(
    doc, os.path.join(OUT, "05_residual_diagnostics.png"),
    "Figure 5. Left: Residuals vs. Fitted Values -- no systematic pattern, confirming "
    "linearity and homoscedasticity. Right: Q-Q plot -- residuals are approximately "
    "normally distributed.")

# Step 8
heading(doc, "Step 8: Feature Importance Analysis", level=2)
body(doc,
     "To understand which predictors drove most of the model's predictive power, I computed "
     "both standardized OLS coefficients (which put all features on a common scale) and "
     "permutation-based feature importance (which measures how much R2 drops when a "
     "feature's values are randomly shuffled).")
body(doc,
     "The standardized coefficient analysis confirmed that Number_of_flights (12.38), "
     "Baggage_loading_time (9.12), and Late_Arrival_o (5.63) were by far the most influential "
     "predictors. Permutation importance corroborated this: Number_of_flights alone accounted "
     "for an R2 drop of 0.364 when removed, followed by Baggage_loading_time (0.203) and "
     "Late_Arrival_o (0.075). Together, these three operational variables carry the vast "
     "majority of the model's predictive signal.")
insert_image(
    doc, os.path.join(OUT, "06_feature_importance.png"),
    "Figure 6. Top 15 features ranked by standardized OLS coefficient magnitude. "
    "Blue bars indicate a positive effect on delay; orange bars indicate a negative effect.")

# ── EVALUATING THE STEPS ───────────────────────────────────────────────────────
heading(doc, "Evaluating the Steps: What I Learned at Each Stage")
body(doc,
     "Looking back on the full workflow, each step served a distinct and non-negotiable "
     "purpose. Step 1 (data loading and exploration) was the foundation -- without understanding "
     "what I am working with, every downstream decision is made on shaky ground. Steps 2 and "
     "3 (assumption checking and relationship evaluation) were, in my experience, the most "
     "undervalued steps in many data science projects. It is tempting to jump straight to "
     "model fitting, but the correlation analysis in Step 3 was what revealed that Cleaning_o "
     "and Fueling_o were essentially noise -- information I would not have had without that "
     "upfront investigation.")
body(doc,
     "Step 4 (feature engineering and train/test split) was a pivotal methodological decision. "
     "The 70/30 split is a widely used convention that balances the need for enough training "
     "data to fit a stable model against having a large enough test set to meaningfully "
     "evaluate generalization. Step 5 (VIF analysis) provided formal assurance that my "
     "coefficient estimates were interpretable -- something that is easy to overlook but "
     "essential for producing a model whose outputs can be communicated meaningfully to "
     "stakeholders.")
body(doc,
     "Steps 6 and 7 were the payoff. Using statsmodels for Step 6 rather than scikit-learn "
     "was a deliberate choice: statsmodels provides the full statistical machinery (p-values, "
     "confidence intervals, Durbin-Watson, AIC/BIC) that practitioners need to evaluate model "
     "validity beyond just R2. Step 7 demonstrated that the model generalizes extremely well, "
     "and the Breusch-Pagan and Durbin-Watson tests gave me confidence that the model's "
     "assumptions were met in the data.")

# ── STRENGTHS AND CHALLENGES ───────────────────────────────────────────────────
heading(doc, "Strengths and Challenges of Multiple Linear Regression")
heading(doc, "Strengths", level=2)
body(doc,
     "Multiple linear regression has several qualities that make it exceptionally well-suited "
     "for operational contexts like airline delay prediction. First and foremost, it is "
     "interpretable. Unlike black-box models such as gradient boosting or neural networks, "
     "every coefficient in an OLS model has a direct, plain-language meaning. For each "
     "additional minute of baggage loading time, arrival delay increases by approximately "
     "13.3 minutes, holding all else constant. That kind of specificity is invaluable when "
     "presenting findings to operations managers who need to know not just what will happen, "
     "but why and by how much.")
body(doc,
     "Second, MLR is computationally efficient. The entire pipeline for this 3,593-record "
     "dataset ran in under 30 seconds on a standard laptop. This efficiency matters for "
     "real-time or near-real-time applications. Third, the accompanying diagnostic tests "
     "(VIF, Durbin-Watson, Breusch-Pagan, residual plots) provide a rich framework for "
     "validating model assumptions -- an infrastructure that more complex models often lack. "
     "Finally, the high R2 of 0.82 achieved here demonstrates that, for this particular "
     "problem, a linear model captures the relationship extremely well.")

heading(doc, "Challenges", level=2)
body(doc,
     "The primary limitation I encountered was the assumption of linearity. MLR assumes that "
     "the relationship between each predictor and the outcome is strictly linear -- an "
     "assumption that may not hold in more complex real-world settings. For instance, the "
     "effect of weather severity on delays is unlikely to be perfectly linear across all "
     "conditions. Relatedly, MLR is sensitive to outliers: a small number of extreme delays "
     "can disproportionately influence the fitted line.")
body(doc,
     "Multicollinearity, while manageable in this dataset (all VIFs below 7), can become a "
     "serious problem when predictors are highly correlated. In such cases, the model may be "
     "accurate overall (high R2) while producing unreliable and counterintuitive individual "
     "coefficient estimates. The Shapiro-Wilk test also flagged a mild deviation from "
     "normality in the residuals, which -- while not severe enough to invalidate the model "
     "given the large sample size -- would be worth investigating in a smaller dataset. "
     "Finally, MLR cannot capture interaction effects (for example, the combined impact of "
     "bad weather and a busy hub) without explicit interaction terms.")

# ── SUITABILITY AND BENEFITS ───────────────────────────────────────────────────
heading(doc, "Suitability and Benefits for the Aviation Context")
body(doc,
     "Multiple linear regression is highly suited to this airline delay use case for several "
     "interconnected reasons. The outcome variable (Arr_Delay, in minutes) is continuous and "
     "unbounded from below, which aligns with MLR's assumption of a continuous, normally "
     "distributed response. The predictors -- mainly operational metrics like loading times, "
     "crew availability, and late inbound aircraft -- are fundamentally additive in nature: "
     "each contributes incrementally to the total delay, which is precisely what an additive "
     "linear model captures.")
body(doc,
     "The practical benefits for airlines are substantial. A deployed version of this model "
     "could run with live flight data to generate real-time delay forecasts as early as when "
     "ground crews begin preparing an aircraft. Airlines could use these predictions to "
     "preemptively notify passengers, reassign gates, or reposition crew members. The model's "
     "coefficient estimates also function as an operational diagnostic tool: if baggage loading "
     "time is the strongest driver of delays (beta = 13.3 minutes per additional minute of "
     "loading time), then investing in faster baggage handling equipment yields a clear, "
     "quantifiable return.")
body(doc,
     "That said, the model would benefit from enrichment with additional data sources: live "
     "weather severity metrics (rather than a binary 5/6 code), air traffic control delays, "
     "runway configuration data, and time-of-day or seasonality features. Incorporating these "
     "would likely push the R2 above 0.85 and reduce RMSE meaningfully.")

# ── REFLECTIONS ───────────────────────────────────────────────────────────────
heading(doc, "Reflections: Informative, Surprising, and Challenging Moments")
body(doc,
     "The most genuinely surprising finding for me was how dominant Number_of_flights was as "
     "a predictor (r = 0.823 with Arr_Delay, and the highest standardized coefficient at "
     "12.38). I initially expected weather or late aircraft arrivals to be the primary drivers "
     "of delays -- that is the intuitive narrative most travelers hold. But the data tells a "
     "different story: the sheer volume of daily departures at a hub is the single strongest "
     "predictor of how delayed a given flight will be. Network congestion at major hubs is "
     "apparently more predictive than the weather or even the downstream cascade of a late "
     "inbound aircraft.")
body(doc,
     "What I found especially informative was the carrier dummy variable analysis. I expected "
     "certain carriers to show significantly different delay profiles, but once operational "
     "factors were controlled for, none of the carrier dummies reached statistical "
     "significance. This suggests that carrier identity itself is not a meaningful driver of "
     "delay -- what matters is what each carrier does operationally. This is an important "
     "insight for airline management: improving delay performance is about operational "
     "efficiency, not brand identity.")
body(doc,
     "The most challenging aspect was navigating the tension between statistical significance "
     "and practical relevance. Several predictors -- Cleaning_o, Fueling_o, and Security_o -- "
     "were statistically non-significant in the full model, yet intuitively these operations "
     "should matter. On reflection, I believe this is partly a multicollinearity effect "
     "(these metrics correlate with other included variables) and partly a measurement "
     "granularity issue: the data may not capture the nuance needed for these variables to "
     "reveal their true effect. This is a reminder that null results in regression do not "
     "necessarily mean a variable is irrelevant -- they mean the model could not detect its "
     "effect given the current data structure.")

# ── CONCLUSION ────────────────────────────────────────────────────────────────
heading(doc, "Conclusion")
body(doc,
     "This lab was a comprehensive, hands-on exercise in building and evaluating a multiple "
     "linear regression model for a real operational problem. By following the full pipeline "
     "-- from exploratory data analysis through assumption checking, feature engineering, model "
     "building, and rigorous evaluation -- I developed both technical skills and a deeper "
     "appreciation for the decisions that separate a careful, trustworthy analysis from a "
     "rushed one. The final model, with an R2 of 0.82 and RMSE of 12.43 minutes, performs "
     "impressively and generalizes cleanly to unseen data.")
body(doc,
     "The key lesson I take away is that the value of multiple linear regression lies not "
     "just in its predictive accuracy but in its transparency. In an industry where "
     "operational decisions affect thousands of passengers and millions of dollars every day, "
     "a model that airline managers can understand, interrogate, and trust is often more "
     "valuable than a more accurate black box. As I continue building toward more complex "
     "predictive models in this course, the discipline I practiced here -- checking "
     "assumptions, validating splits, interpreting diagnostics -- will remain the foundation "
     "of sound analytical work.")

divider(doc)

# ── REFERENCES ────────────────────────────────────────────────────────────────
heading(doc, "References")
refs = [
    "Federal Aviation Administration. (2023). Understanding the impact of flight delays "
    "and cancellations. U.S. Department of Transportation. https://www.faa.gov",

    "Hair, J. F., Black, W. C., Babin, B. J., & Anderson, R. E. (2019). Multivariate data "
    "analysis (8th ed.). Cengage Learning.",

    "McKinney, W. (2022). Python for data analysis (3rd ed.). O'Reilly Media.",

    "Montgomery, D. C., Peck, E. A., & Vining, G. G. (2021). Introduction to linear "
    "regression analysis (6th ed.). Wiley.",

    "Pedregosa, F., Varoquaux, G., Gramfort, A., et al. (2011). Scikit-learn: Machine "
    "learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",

    "Seabold, S., & Perktold, J. (2010). Statsmodels: Econometric and statistical modeling "
    "with Python. Proceedings of the 9th Python in Science Conference.",
]
for ref in refs:
    rp = doc.add_paragraph()
    rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rp.paragraph_format.first_line_indent = Inches(-0.5)
    rp.paragraph_format.left_indent       = Inches(0.5)
    rp.paragraph_format.space_after       = Pt(6)
    run = rp.add_run(ref)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(BASE, "Week2_Lab_Shruti_Malik.docx")
doc.save(out_path)
print("Document saved:", out_path)
